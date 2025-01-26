import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
from fpdf import FPDF
from docx2pdf import convert
import tempfile
import os

def format_description(description):
    """Converts multi-sentence descriptions into bullet points."""
    points = [line.strip() for line in description.split("\n") if line.strip()]
    return points

def create_word_doc(resume_data):
    """Creates a Word document from resume data."""
    doc = Document()

    doc.add_heading(resume_data["Name"], level=1)
    doc.add_paragraph(f"Email: {resume_data['Email']}")
    doc.add_paragraph(f"Phone: {resume_data['Phone']}")
    doc.add_paragraph(f"LinkedIn: {resume_data['LinkedIn']}")
    doc.add_paragraph(f"GitHub: {resume_data['GitHub']}")

    doc.add_heading("Professional Summary", level=2)
    doc.add_paragraph(resume_data["Summary"])

    doc.add_heading("Work Experience", level=2)
    for exp in resume_data["Work Experience"]:
        doc.add_paragraph(f"{exp['Role']} at {exp['Company']} ({exp['Duration']})")
        for point in exp['Description']:
            doc.add_paragraph(f"- {point}", style="List Bullet")

    doc.add_heading("Education", level=2)
    for edu in resume_data["Education"]:
        doc.add_paragraph(f"{edu['Degree']} from {edu['Institution']} ({edu['Year']})")

    doc.add_heading("Skills", level=2)
    doc.add_paragraph(resume_data["Skills"])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_pdf_from_docx(word_doc_buffer):
    """Converts a Word document buffer to a PDF using docx2pdf."""
    with tempfile.TemporaryDirectory() as temp_dir:
        word_path = os.path.join(temp_dir, "resume.docx")
        pdf_path = os.path.join(temp_dir, "resume.pdf")

        # Save Word document buffer to temporary file
        with open(word_path, "wb") as f:
            f.write(word_doc_buffer.getbuffer())

        # Convert Word document to PDF
        convert(word_path, pdf_path)

        # Read the generated PDF into a buffer
        with open(pdf_path, "rb") as f:
            pdf_buffer = BytesIO(f.read())

    return pdf_buffer

def main():
    st.title("Resume Builder")

    st.sidebar.title("Resume Sections")

    # Basic Information
    st.sidebar.header("Personal Details")
    name = st.sidebar.text_input("Name")
    email = st.sidebar.text_input("Email")
    phone = st.sidebar.text_input("Phone")
    linkedin = st.sidebar.text_input("LinkedIn Profile")
    github = st.sidebar.text_input("GitHub Profile")

    # Summary Section
    st.sidebar.header("Professional Summary")
    summary = st.sidebar.text_area("Summary", "Write a brief summary about yourself.")

    # Work Experience Section
    st.sidebar.header("Work Experience")
    experiences = []
    for i in range(1, 4):
        st.sidebar.subheader(f"Experience {i}")
        role = st.sidebar.text_input(f"Role {i}", key=f"role_{i}")
        company = st.sidebar.text_input(f"Company {i}", key=f"company_{i}")
        duration = st.sidebar.text_input(f"Duration {i}", key=f"duration_{i}")
        description = st.sidebar.text_area(f"Description {i} (use blank lines to separate points)", key=f"description_{i}")

        if role and company and duration and description:
            formatted_description = format_description(description)
            experiences.append({
                "Role": role,
                "Company": company,
                "Duration": duration,
                "Description": formatted_description,
            })

    # Education Section
    st.sidebar.header("Education")
    education = []
    for i in range(1, 3):
        st.sidebar.subheader(f"Education {i}")
        degree = st.sidebar.text_input(f"Degree {i}", key=f"degree_{i}")
        institution = st.sidebar.text_input(f"Institution {i}", key=f"institution_{i}")
        year = st.sidebar.text_input(f"Year {i}", key=f"year_{i}")

        if degree and institution and year:
            education.append({
                "Degree": degree,
                "Institution": institution,
                "Year": year,
            })

    # Skills Section
    st.sidebar.header("Skills")
    skills = st.sidebar.text_area("List your skills (separated by commas)")

    # Display Resume
    st.header("Preview Your Resume")

    st.write("### Personal Details")
    st.write(f"**Name:** {name}")
    st.write(f"**Email:** {email}")
    st.write(f"**Phone:** {phone}")
    st.write(f"**LinkedIn:** {linkedin}")
    st.write(f"**GitHub:** {github}")

    st.write("### Professional Summary")
    st.write(summary)

    st.write("### Work Experience")
    for exp in experiences:
        st.write(f"**{exp['Role']}** at {exp['Company']} ({exp['Duration']})")
        for point in exp['Description']:
            st.write(f"- {point}")

    st.write("### Education")
    for edu in education:
        st.write(f"**{edu['Degree']}** from {edu['Institution']} ({edu['Year']})")

    st.write("### Skills")
    st.write(skills)

    # Generate Resume Data
    resume_data = {
        "Name": name,
        "Email": email,
        "Phone": phone,
        "LinkedIn": linkedin,
        "GitHub": github,
        "Summary": summary,
        "Work Experience": experiences,
        "Education": education,
        "Skills": skills,
    }

    # Download Resume as Word
    if st.button("Download Resume as Word Document"):
        word_doc = create_word_doc(resume_data)
        st.download_button(
            label="Download Word Document",
            data=word_doc,
            file_name="resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # Download Resume as PDF
    if st.button("Download Resume as PDF"):
        word_doc = create_word_doc(resume_data)
        pdf_doc = create_pdf_from_docx(word_doc)
        st.download_button(
            label="Download PDF",
            data=pdf_doc,
            file_name="resume.pdf",
            mime="application/pdf"
        )

if __name__ == "__main__":
    main()
